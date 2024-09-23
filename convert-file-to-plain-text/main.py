import PyPDF2
import docx


def extract_text_from_pdf(pdf_file: str) -> str:
    with open(pdf_file, 'rb') as pdf:
        reader = PyPDF2.PdfReader(pdf)
        pdf_text = []

        for page in reader.pages:
            content = page.extract_text()
            if content:
                pdf_text.append(content)

        return '\n'.join(pdf_text)
    
def extract_text_from_docx(docx_file: str) -> str:
    # Membuka file .docx
    doc = docx.Document(docx_file)
    paragraph_count = len(doc.paragraphs)
    full_text = []
    for ind, paragraph in enumerate(doc.paragraphs):
        if ind + 1 < paragraph_count:
            if doc.paragraphs[ind + 1].style.name.startswith('Heading'):
                print('aaa')
                full_text.append(paragraph.text + '\n\n')
            else:
                full_text.append(paragraph.text)
        else:
            full_text.append(paragraph.text)
    
    # Menggabungkan semua paragraf dengan newline
    return '\n'.join(full_text)

def save_text_to_file(text: str, output_file: str):
    # Replace actual newlines with the string "\n" for literal representation
    # text_with_escaped_newlines = text.replace('\n', '\\n')

    # Save the text with literal \n into a .txt file
    with open(output_file, 'w') as file:
        file.write(text)

def main():
    extracted_text = extract_text_from_docx('Maria Stefany-resume New.docx')

    # Save the extracted text to a .txt file
    save_text_to_file(extracted_text, 'output.txt')


main()
